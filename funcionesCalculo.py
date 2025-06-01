
# Libreria de funciones usadas en los cálculos 

# Listado de funciones (actualizar según se añaden funciones)

    # crea_directorio(), genera el directorio de trabajo para alojar los archivos de proyecto
    # datos_terraplen, importa los datos de la carga del terraplen y de la malla de cálculo
    # datos_terreno, importa de una hoja excel los datos del terreno
    # tension_terraplen, calcula las tensiones, (x,z,xz) del terreno

    # asiento_elastico, calcula el asiento elástico del terreno
    # asiento_consolidación, calcula el asiento por consolidación

    # parametro_terreno, obtiene cualquier parámetro del terreno en función de la profundidad
    # n_freatico, calcula si a una profundidad existe nivel freático
    # insertar_valor inserta un valor en una lista y la ordena con el nuevo valor
        # es función auxiliar de presion_total
    # obtener_maximo_menor calculamos el valor de la cota inferior del estrato superior más cercano a la cota que se introduce
        # es función  auxiliar de presion_total
    # presion_total, calcula la presión total en un punto del terreno
    # insertar_valor, inserta un valor de forma ordenada dentro de una lista
    # resistencia_MC, calcula el valor de la resistencia al corte
    # guardar_docx_datos, se guarda un resumen de los datos y resultados en formato word
    # guardar_xlxs_tensiones, guerda en formato excel los resultados de los cálculos de las tensiones creados
        # por la carga del terraplén, en x, z, xz
    # guardar_xlxs_asientos, guarda en excel los resultados de los cálculos de los asientos
    
    # ploteado_tensiones_normales, obtiene la grafica de tensiones normales totales
    # graficos_tensiones, guarda los graficos de las tensiones producidas por el terraplén, en contínuo e isolínea
    # graficos_asientos, guarda la gráfica de los asientos del terraplén en el ancho de banda




# librerias 
import numpy as np
import matplotlib.pyplot as plt
import openpyxl
from docx import Document
from docx.shared import Cm, Pt
import os
from datetime import datetime


# Grupo de funciones


def crea_directorio():
    # creación del directorio de trabajo para guardar resultados
    now = datetime.now()
    directorio=(str(now.year)+str(now.month)+str(now.day)+str(now.hour)+str(now.minute)+str(now.second))
    os.mkdir(directorio)
    return directorio


def datos_terraplen():
    # importacion de los datos del terraplen
    # desde la hoja excel
    libro = openpyxl.load_workbook('datos_terraplen.xlsx')
    hoja = libro.active 

    # datos del terraplen 
    a=hoja.cell(row=2, column=1).value # ancho del derrame del terraplen
    b=hoja.cell(row=2, column=2).value # semi ancho del terraplén
    h=hoja.cell(row=2, column=3).value # altura del terraplén
    pe=hoja.cell(row=2, column=4).value # peso especifico del terraplen

    # datos del mallado de cálculo
    ax=hoja.cell(row=2, column=5).value # ancho de banda
    incrx=hoja.cell(row=2, column=6).value # incremento de x
    incrz=hoja.cell(row=2, column=7).value # incremento de z
    
    return a,b,h,pe*h,ax,incrx,incrz


def datos_terreno():

    libro = openpyxl.load_workbook('datos_terreno.xlsx')
    hoja = libro.active

    # importacion de variables del terreno
    
    # geometría de las capas
    espesor=[]
    cotas=[]
    # valores físicos
    pe_seco=[]
    pe_saturado=[]
    # valores elásticos
    E=[]
    poisson=[]
    tipo_datos=[]
    # valores de resistencia
    cohesion=[]
    fi=[]
    # valores de compresibilidad por consolidacion
    cc=[]
    e0=[]

    # tipo de calculo a realizar 
    # (E,e) de tipo elástico
    # (C,c) de tipo consolidación
    tipo_calculo=[]
   
    for row in hoja.iter_rows():
        espesor.append(row[0].value)
        espesor[0]=0
        az=sum(espesor) # vector de espesores
        nivel_freatico=hoja.cell(row=2, column=2).value
        pe_seco.append(row[2].value)
        pe_seco[0]=0
        pe_saturado.append(row[3].value)
        pe_saturado[0]=0
        E.append(row[4].value)
        E[0]=0
        poisson.append(row[5].value)
        poisson[0]=0
        cohesion.append(row[6].value)
        cohesion[0]=0
        fi.append(row[7].value)
        fi[0]=0
        cc.append(row[8].value)
        cc[0]=0
        e0.append(row[9].value)
        e0[0]=0
        tipo_calculo.append(row[10].value)
        tipo_calculo[0]=0


    for i in np.arange(len(espesor)):
        cotas.append(sum(espesor[0:i+1]))
    
    # se toman los datos de la cabecera
    for filas in hoja.iter_cols():
        tipo_datos.append(filas[0].value)

    return espesor,cotas,az,nivel_freatico,pe_seco,pe_saturado,E,poisson,cohesion,fi,cc,e0,tipo_datos,tipo_calculo


def tension_terraplen(a,b,q,x,z):
    # a es la base del derrame del terraplén
    # b es el ancho de la base de medio terraplén
    # q es la carga del terraplen en la zona central
    # pe peso especifico del relleno de terraplén
    # x,z son las coordenadas del punto en el que queremos calcular las tensiones y deformaciones
    # calculo de las tensiones bajo un terraplen
    # los valores de las tensiones se calculan desde el borde exterior izquierdo hacia dentro del terraplén

    # cálculos intermedios
    beta_a=np.arctan((b-x)/z)+np.arctan((x-a)/z)
    beta_b=np.arctan((x-b)/z)+np.arctan((2*b-x-a)/z)
    alfa_a=np.arctan((a-x)/z)+np.arctan(x/z)
    alfa_b=np.arctan((a-2*b+x)/z)+np.arctan((2*b-x)/z)

    r_oa=np.sqrt(np.power(x,2)+np.power(z,2))
    r_ob=np.sqrt(np.power(2*b-x,2)+np.power(z,2))
    r_1a=np.sqrt(np.power(x-a,2)+np.power(z,2))
    r_1b=np.sqrt(np.power(2*b-x-a,2)+np.power(z,2))
    r_22=np.power(b-x,2)+np.power(z,2)

    # cálculos de los incrementos de tensiones en x,z,xz
    tensionz=(q/np.pi)*((beta_a+x*alfa_a/a-z*(x-b)/r_22)+(beta_b+(2*b-x)*alfa_b/a-z*(b-x)/r_22))
    tensionx=(q/np.pi)*((beta_a+x*alfa_a/a+z*(x-b)/r_22+2*z*np.log(r_1a/r_oa)/a)+beta_b+alfa_b*(2*b-x)/a+z*(b-x)/r_22+2*z*np.log(r_1b/r_ob)/a)
    tensionxz=-(q/np.pi)*(z*(alfa_a+alfa_b)/a-2*np.power(z,2)/r_22)

    return tensionz,tensionx,tensionxz


def asiento_elastico(cotas,z,hi,E,poisson,tensionx,tensionz):
    # calculo de aientos para un estado bidimensional de tensiones
    # se calcula para un valor de z determinado, para más valores incluirlo en un bucle
    # los valores en cm

    # versión teniendo en cuenta el efecto tridimensional
    asiento=-hi*(tensionz-poisson[parametro_terreno(cotas,z)]*tensionx)/E[parametro_terreno(cotas,z)]

    # se pone el valor negativo suponiendo el valor vertical hacia abajo

    return asiento

def asiento_consolidacion(hi,cc,e0,t_efectiva,tensionz,cotas,x,z):
    # calculo de aientos por consolidación
    # parametros de entrada 
    # e0, Cc y los incrementos de tensiones por la sobre carga y las tensiones efectivas naturales
    # se debe de considerar las tres condiciones de OCR, pero por ahora solo el caso de OCR=1
    asiento=-(hi*cc[parametro_terreno(cotas,z)]/(1+e0[parametro_terreno(cotas,z)]))*np.log10((t_efectiva+tensionz)/t_efectiva)
    
    return asiento

def parametro_terreno(cotas,zt):
    # esta funcion localiza el parámetro del terreno para una z determinada
    # se parte de los valores discretos tomados de la excel de datos_terreno.xlsx
    # sirve para cualquier parámetro del terreno

    for z in np.arange(len(cotas)-1):
        cota_superior=cotas[z]
        cota_inferior=cotas[z+1]
        if zt>=cota_superior and zt<=cota_inferior:
            break
    return z+1



def n_freatico(nivel_freatico,z):
    # función para calcular el nivel freático de forma continua
    # indica si un valor del terreno está o no bajo el nf
    # devuelve el valor de la altura de columna o un 0 según se esté o no en la zona del nivel freático
    if z>=nivel_freatico:
        return (z-nivel_freatico)
    else:
        return 0


# funciones auxiliares para el cálculo de la presion total

def insertar_valor(lista, valor):
    # funcion para insertar la cota en la lista de cotas y ordenarla
    lista_mod=[] 
    lista_mod=lista.copy() # se evita alterar la lista original
    lista_mod.append(valor)  # Agrega el valor a la lista
    lista_mod.sort()         # Ordena la lista en orden ascendente
    return lista_mod



def obtener_maximo_menor(lista, valor):
    # calculamos el valor de la cota inferior del estrato superior 
    # más cercano a la cota que se introduce
    maximo_menor = max(filter(lambda x: x <= valor, lista))
    return maximo_menor



def presion_total(cotas,valor_nf,pe_saturado,pe_seco,valor_cota):
    
    # cotas es la lista de cotas
    # pe_saturado, lista de pesos específicos saturados
    # pe_seco.lista de pesos específicos secos
    # nf es la profundidad del nivel freático
    # z es la profundidad a la que se va a calcular la presión total
    
    
    lista_cotas = cotas.copy() # se copia la lista de las cotas para no alterarla

    #lista ordenada contiene la lista de las cotas y la cota hasta la que se quiere calcular las tensiones
    # en caso de que ya esté el valor no se introduce de nuevo

    if valor_nf not in lista_cotas: 
        lista_valores = insertar_valor(lista_cotas, valor_nf)
    else:
        lista_valores=lista_cotas

    
    # resultado es el valor de la cota inmediatamente anterior a la que queremos calcular
    resultado = obtener_maximo_menor(lista_valores, valor_cota)
    
    
    # calculo del estrato del fondo
    peso_saturado=pe_saturado[parametro_terreno(cotas,valor_cota)] # peso especifico saturado
    peso_seco=pe_seco[parametro_terreno(cotas, valor_cota)] # peso especifico seco
    
    # cálculo de las presiones totales y efectivas, inicio de valores
    peso=peso_seco if valor_cota<=valor_nf else peso_saturado
    presion_total=(valor_cota-resultado)*peso
    
    # resto de estratos
    for j in range(lista_valores.index(resultado),0,-1):
        espesor=lista_valores[j]-lista_valores[j-1] # espesor del estrato
        peso_saturado=pe_saturado[parametro_terreno(cotas,lista_valores[j])] # peso especifico saturado
        peso_seco=pe_seco[parametro_terreno(cotas,lista_valores[j])] # peso especifico seco
        posicion=lista_valores[j] # indica la cota del final del nivel
    
        # sumatoria de las presiones
        if (posicion<=valor_nf):
            peso=peso_seco
        else:
            peso=peso_saturado
    
        presion_total=presion_total+espesor*peso
    
    return presion_total



def resistencia_MC(cotas,valor_presion,cohesion,fi,z):
    # calculo de la resisencia al corte en un punto del terreno
    # el valor de la presion puede ser en totales o en efectivas

    # cálculo de la cohesión y el ángulo de rozamiento
    cohesion_ter=cohesion[parametro_terreno(cotas,z)]
    fi_ter=fi[parametro_terreno(cotas,z)]
    
    # valor de la resistencia al corte 
    res_corte=cohesion_ter+valor_presion*np.tan(np.deg2rad(fi_ter))

    return res_corte



def guardar_docx_datos(a,b,h,q,ax,incrx,az,incrz,directorio,nivel_freatico,asiento_max):
    # creación de un informe con los datos de entrada y los resultados en formato  .docx
    # 
    
    document = Document()
    # Añadimos un titulo al documento, a nivel 0
    document.add_heading('Datos del modelo geométrico', 0)
    # Añadimos un párrafo
    p = document.add_paragraph('En este documento se incluyen los datos de la carga y características del terreno, ')
    p.add_run('Así como los resultados de las tensiones y asientos ')
    
    # Datos iniciales
    # Características del terraplén
    document.add_heading('Datos del terraplen ', level=1)
    document.add_paragraph("Valor del ancho del derrame {0:0.2f} m".format(a), style='List Bullet')
    document.add_paragraph("Valor del ancho de la semibase {0:0.2f} m".format(b), style='List Bullet')
    document.add_paragraph("Valor de la altura del terraplén {0:0.2f} m".format(h), style='List Bullet')
    document.add_paragraph("Peso específico del relleno de terraplén {0:0.2f} kN/m3".format(q/h), style='List Bullet')
    document.add_paragraph("Valor de la carga de tierras  {0:0.2f} kN/m2".format(q), style='List Bullet')
    document.add_paragraph("Profundidad del nivel freático  {0:0.2f} m".format(nivel_freatico), style='List Bullet')

    # Características del mallado de cálculo
    document.add_heading('Datos del mallado ', level=1)
    document.add_paragraph("Ancho de banda {0:0.2f} m".format(ax), style='List Bullet')
    document.add_paragraph("incremento en x= {0:0.2f} m".format(incrx), style='List Bullet')
    document.add_paragraph("Profundidad de cálculo {0:0.2f} m".format(az), style='List Bullet')
    document.add_paragraph("incremento en z= {0:0.2f} m".format(incrz), style='List Bullet')


    # Tabla de datos del terreno
    # a variar según se incrementa el numero de parámetros a tener en cuenta en los cálculos 

    # Añadimos un titulo al documento, a nivel 0
    document.add_heading('Datos del modelo geotécnico ', 0)
    
    # Ruta del archivo Excel
    excel_file = 'datos_terreno.xlsx'

    # Abre el archivo Excel
    workbook = openpyxl.load_workbook(excel_file)

    # Selecciona la hoja de Excel
    sheet = workbook['Hoja1']

    # Crea un nuevo documento de Word
    #doc = Document()

    # Crea una tabla en el documento de Word
    table = document.add_table(rows=0, cols=sheet.max_column)
    table.autofit = False
    # Establecer el tamaño de fuente para todas las celdas de la tabla
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(9)


    # Obtén los datos de Excel y añádelos a la tabla de Word
    for row in sheet.iter_rows():
        table_row = table.add_row().cells
        for i, cell in enumerate(row):
            table_row[i].text = str(cell.value)


    # Ajusta el ancho de las celdas para que se adecuen al contenido
    for column in table.columns:
        for cell in column.cells:
            cell.width=Cm(1.5)


# introducimos dos retorno de carro después de la tabla

    paragraph = document.add_paragraph("\n\n")

# Añadimos un titulo al documento, a nivel 0
    document.add_heading('Resultados de los cálculos', 0)


    document.add_heading('Cálculos realizados ', level=1)
    document.add_paragraph('Asientos bajo el terraplén [cm]', style='List Number')
    document.add_paragraph('Tensiones efectivas en el terreno [kN/m2]', style='List Number')
    document.add_paragraph('Tensiones totales en el terreno [kN/m2]', style='List Number')
    document.add_paragraph('Incremento de tensiones en x [kN/m2]', style='List Number')
    document.add_paragraph('Incremento de tensiones en z [kN/m2]', style='List Number')
    document.add_paragraph('Incremento de tensiones en xz [kN/m2]', style='List Number')
    

    document.add_heading('Imágenes de los resultados ', level=1)

    # imágenes de resultados de los cálculos
    # se localizan las imágenes del directorio que sean *.png y se meten en unna lista
    listado=os.listdir(directorio)
    imagenes = []
    for fichero in listado:
        if os.path.isfile(os.path.join(directorio, fichero)) and fichero.endswith('.png'):
            imagenes.append(fichero)
    imagenes=sorted(imagenes) # ordenacion de la lista de imágenes

    # carga de la imagen de asientos
    document.add_heading('Asientos', 3) 
    document.add_picture(directorio+'/'+imagenes[-1],width=Cm(12))
    
    # representación del asiento máximo
    numero_formateado = "{:.4f}".format(asiento_max)
    paragraph = document.add_paragraph("Asiento máximo "+numero_formateado+" m")

    
    # carga del resto de imágenes
    for index in np.arange(0,len(imagenes)-1):
        document.add_heading(imagenes[index][0:len(imagenes[index])-4], 3) 
        document.add_picture(directorio+'/'+imagenes[index],width=Cm(12))

    # Guardado del archivo con la información
    document.save(directorio+'/'+'Informe.docx')




def guardar_xlsx_tensiones(xcoord,zcoord,array_datos,directorio,nombre_archivo):

    # guardado en excel de los resultados de los calculos de una matriz de datos
    wb = openpyxl.Workbook()
    hoja = wb.active
    for fila in np.arange(zcoord.size):
        hoja.cell(row=fila+2, column=1, value=zcoord[fila])
        for columna in np.arange(xcoord.size):
            hoja.cell(row=fila+2, column=columna+2, value=array_datos[fila,columna])

     # valores de x
    for columna in np.arange(xcoord.size):
        hoja.cell(row=1, column=columna+2, value=xcoord[columna])

    wb.save(directorio+'/'+nombre_archivo+'.xlsx')
    



def guardar_xlsx_asientos(xcoord,array_datos,directorio,nombre_archivo):
    
    # guardado en excel de los resultados de los calculos de una matriz de datos
    wb = openpyxl.Workbook()
    hoja = wb.active
    # valores del vector
    for columna in np.arange(xcoord.size):
        hoja.cell(row=2, column=columna+1, value=array_datos[columna])
     # valores de x
    for columna in np.arange(xcoord.size):
        hoja.cell(row=1, column=columna+1, value=xcoord[columna])

    # se crea un directorio para calculo realizado

    wb.save(directorio+'/'+nombre_archivo+'.xlsx')




def graficos_tensiones(xcoord,zcoord,tension,directorio,titulo,tipo,a,b,h):
    # representacion gráfica de los resultados en modo continuo o en isolinea
    X, Z = np.meshgrid(xcoord, zcoord)
    fig, ax = plt.subplots()

    # grafica del terraplen
    x = [-b,-(b-a),0,b-a,b]
    y=[0,h,h,h,0]
    plt.plot(x,y)

    # grafias de tipo contorno o isolineas
    if tipo=='isolinea': # tipo isolinea
        # suelo del terraplen
        x = [min(xcoord),max(xcoord)]
        y=[0,0]
        plt.plot(x,y)
        # grafica de isolinea
        curvas=plt.contour(X,-Z,tension,10)
        plt.clabel(curvas, inline=1,fmt='%2.1f',fontsize=8) # etiquetas
    else: # tipo contorno
        curvas=plt.contourf(X,-Z,tension,10)
        cbar=plt.colorbar(curvas) # leyenda del grafico
        # Invertir el orden de la leyenda de valores
   

    ax.set_aspect('equal', adjustable='box')
    #ax.invert_yaxis()
    plt.xlabel("x [m]")
    plt.ylabel("z [m]")
    ax.set_title(titulo+' [kN/m2]')
    plt.savefig(directorio+'/'+titulo+tipo+".png") # guardado de la imagen
    #plt.show()




def grafico_asientos(xcoord,asiento,directorio,titulo):
    
    fig, ax = plt.subplots()

    # grafica del terraplén desactivada
    #x = [-b,-(b-a),0,b-a,b]
    #y=[0,h,h,h,0]
    #plt.plot(x,y)

    # linea horizontal
    #x = [min(xcoord),max(xcoord)]
    #y=[0,0]
    #plt.plot(x,y)

    # trazado de los asientos
    plt.plot(xcoord,np.multiply(asiento,1),marker='*',color='darkred',linestyle='-')
    ax.set_aspect('auto', adjustable='box')
    #ax.invert_yaxis()
    plt.xlabel("x [m]")
    plt.ylabel("asiento [m]")
    ax.set_title(titulo)
    #ax.set_aspect('equal', adjustable='box')
    plt.savefig(directorio+'/'+titulo+".png") # guardado de la imagen
    #plt.show()


